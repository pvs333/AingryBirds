from docx import Document
from docx.shared import Inches
import pandas as pd
import os

doc = Document()
doc.add_heading('Angry Birds AI/ML Project Results', 0)

doc.add_heading('Key Metrics', level=1)
metrics = pd.read_csv('ML/report/key_metrics.csv')
table = doc.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Metric'
hdr_cells[1].text = 'Value'
for _, row in metrics.iterrows():
    cells = table.add_row().cells
    cells[0].text = str(row['Metric'])
    cells[1].text = f"{row['Value']:.3f}" if isinstance(row['Value'], float) else str(row['Value'])

doc.add_heading('Sample Last 10 Episodes', level=1)
episodes = pd.read_csv('ML/report/sample_last_10_episodes.csv')
table2 = doc.add_table(rows=1, cols=len(episodes.columns))
for i, col in enumerate(episodes.columns):
    # Rename columns for clarity in the DOCX
    if col == 'With Heuristic Data Success':
        table2.rows[0].cells[i].text = 'With Heuristic Data Success Rate'
    elif col == 'Without Heuristic Data Success':
        table2.rows[0].cells[i].text = 'Without Heuristic Data Success Rate'
    else:
        table2.rows[0].cells[i].text = col
for _, row in episodes.iterrows():
    cells = table2.add_row().cells
    for i, col in enumerate(episodes.columns):
        val = row[col]
        if isinstance(val, float):
            cells[i].text = f"{val:.3f}"
        else:
            cells[i].text = str(val)

doc.add_page_break()

doc.save('ML/report/AngryBirds_AI_ML_Results.docx')
print('✅ DOCX report generated!')
